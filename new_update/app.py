from flask import Flask, request, jsonify, send_file, send_from_directory, render_template_string, redirect, url_for
from flask_cors import CORS
import os
from config import Config
import tempfile
import io
import uuid
from pathlib import Path
import logging
from datetime import datetime
import mammoth
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
import re

# Importations pour le traitement des documents
try:
    from docx import Document
    from docx2txt import process as docx2txt_process
except ImportError:
    print("Installez les dépendances: pip install python-docx docx2txt")

try:
    import mammoth
except ImportError:
    print("Installez mammoth: pip install mammoth")

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.units import inch
except ImportError:
    print("Installez reportlab: pip install reportlab")

try:
    import markdown2
except ImportError:
    print("Installez markdown2: pip install markdown2")

# Configuration de l'application Flask
app = Flask(__name__, 
            static_folder='static',
            static_url_path='')
app.config.from_object(Config)
CORS(app)  # Permet les requêtes cross-origin

# Configuration pour servir les fichiers statiques (CSS, JS, images)
@app.route('/<path:path>')
def serve_static(path):
    if path.startswith('api/'):
        return jsonify({'error': 'Not found'}), 404
    return send_from_directory('static', path)

# Routes principales
@app.route('/')
def home():
    return send_from_directory('static', 'index.html')

@app.route('/editor')
def editor():
    return send_from_directory('static', 'editor.html')

# Configuration
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024  # 10MB max
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'

# Créer les dossiers s'ils n'existent pas
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# Configuration du logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Classe pour traiter les documents Word"""
    
    @staticmethod
    def convert_docx_to_html(file_path):
        """
        Convertit un fichier DOCX en HTML en préservant la mise en forme
        """
        try:
            # Utiliser mammoth pour la conversion
            with open(file_path, "rb") as docx_file:
                result = mammoth.convert_to_html(docx_file)
                html = result.value
                
                # Nettoyer et formater le HTML
                soup = BeautifulSoup(html, 'html.parser')
                
                # Créer la balise style avec les styles CSS
                style_content = """
                    body { 
                        font-family: Arial, sans-serif; 
                        line-height: 1.6; 
                        max-width: 8.5in;
                        margin: 0 auto;
                        padding: 1in;
                    }
                    h1, h2, h3, h4, h5, h6 { 
                        color: #2c3e50; 
                        margin-top: 1.5em;
                    }
                    p { 
                        margin: 0 0 1em 0;
                    }
                    table {
                        border-collapse: collapse;
                        width: 100%;
                        margin: 1em 0;
                    }
                    table, th, td {
                        border: 1px solid #ddd;
                    }
                    th, td {
                        padding: 8px;
                        text-align: left;
                    }
                """
                
                # Créer une nouvelle balise style
                style_tag = soup.new_tag('style')
                style_tag.string = style_content
                
                # Vérifier si la balise head existe, sinon la créer
                if soup.head is None:
                    html_tag = soup.find('html')
                    if not html_tag:
                        html_tag = soup.new_tag('html')
                        soup.append(html_tag)
                    head_tag = soup.new_tag('head')
                    html_tag.insert(0, head_tag)
                
                # Ajouter le style au head
                soup.head.append(style_tag)
                
                return str(soup)
                
        except Exception as e:
            app.logger.error(f"Erreur lors de la conversion DOCX vers HTML: {str(e)}")
            raise
    
    @staticmethod
    def convert_html_to_docx(html_content, output_path):
        """
        Convertit du HTML en document Word (DOCX) avec mise en forme
        """
        try:
            # Créer un nouveau document Word
            doc = Document()
            
            # Configurer les styles par défaut
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(11)
            
            # Parser le HTML avec BeautifulSoup
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Supprimer les balises script et style
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Convertir les balises HTML en éléments Word
            for element in soup.find_all(True):
                if element.name == 'h1':
                    p = doc.add_heading(level=1)
                    p.add_run(element.get_text())
                    p.style = 'Heading 1'
                elif element.name == 'h2':
                    p = doc.add_heading(level=2)
                    p.add_run(element.get_text())
                    p.style = 'Heading 2'
                elif element.name == 'p':
                    p = doc.add_paragraph()
                    p.add_run(element.get_text())
                elif element.name in ['ul', 'ol']:
                    for li in element.find_all('li'):
                        p = doc.add_paragraph(style='List Bullet' if element.name == 'ul' else 'List Number')
                        p.add_run(li.get_text())
                elif element.name == 'table':
                    # Créer un tableau dans le document
                    rows = element.find_all('tr')
                    if rows:
                        # Déterminer le nombre de colonnes
                        cols = max(len(row.find_all(['th', 'td'])) for row in rows)
                        table = doc.add_table(rows=0, cols=cols)
                        table.style = 'Table Grid'
                        
                        # Ajouter les en-têtes
                        headers = rows[0].find_all(['th', 'td'])
                        if headers:
                            row_cells = table.add_row().cells
                            for i, header in enumerate(headers):
                                row_cells[i].text = header.get_text()
                            
                            # Ajouter les lignes de données
                            for row in rows[1:]:
                                cells = row.find_all('td')
                                if cells:
                                    row_cells = table.add_row().cells
                                    for i, cell in enumerate(cells):
                                        if i < len(row_cells):
                                            row_cells[i].text = cell.get_text()
            
            # Enregistrer le document
            doc.save(output_path)
            return True
            
        except Exception as e:
            app.logger.error(f"Erreur lors de la conversion HTML vers DOCX: {str(e)}")
            raise
    
    @staticmethod
    def extract_text_from_docx(file_path):
        """Extrait le texte d'un fichier .docx"""
        try:
            doc = Document(file_path)
            text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            return '\n'.join(text)
        except Exception as e:
            logger.error(f"Erreur lors de l'extraction du texte: {e}")
            # Fallback avec docx2txt
            try:
                return docx2txt_process(file_path)
            except:
                return ""
    
    @staticmethod
    def extract_text_from_doc(file_path):
        """Extrait le texte d'un fichier .doc avec mammoth"""
        try:
            with open(file_path, "rb") as docx_file:
                result = mammoth.extract_raw_text(docx_file)
                return result.value
        except Exception as e:
            logger.error(f"Erreur lors de l'extraction du fichier .doc: {e}")
            return ""
    
    @staticmethod
    def convert_to_pdf(text, output_path):
        """Convertit le texte en PDF"""
        try:
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Diviser le texte en paragraphes
            paragraphs = text.split('\n')
            
            for para_text in paragraphs:
                if para_text.strip():
                    para = Paragraph(para_text, styles['Normal'])
                    story.append(para)
                    story.append(Spacer(1, 12))
            
            doc.build(story)
            return True
        except Exception as e:
            logger.error(f"Erreur lors de la conversion en PDF: {e}")
            return False
    
    @staticmethod
    def convert_to_html(text, output_path):
        """Convertit le texte en HTML"""
        try:
            html_content = f"""
<!DOCTYPE html>
<html lang="fr">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Document converti</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            line-height: 1.6;
            margin: 40px;
            color: #333;
        }}
        p {{
            margin-bottom: 15px;
        }}
    </style>
</head>
<body>
    <h1>Document converti</h1>
    <div class="content">
"""
            
            # Convertir chaque ligne en paragraphe HTML
            paragraphs = text.split('\n')
            for para in paragraphs:
                if para.strip():
                    html_content += f"        <p>{para}</p>\n"
            
            html_content += """
    </div>
</body>
</html>
"""
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            return True
        except Exception as e:
            logger.error(f"Erreur lors de la conversion en HTML: {e}")
            return False
    
    @staticmethod
    def convert_to_markdown(text, output_path):
        """Convertit le texte en Markdown"""
        try:
            markdown_content = "# Document converti\n\n"
            
            paragraphs = text.split('\n')
            for para in paragraphs:
                if para.strip():
                    markdown_content += f"{para}\n\n"
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            return True
        except Exception as e:
            logger.error(f"Erreur lors de la conversion en Markdown: {e}")
            return False
    
    @staticmethod
    def convert_to_txt(text, output_path):
        """Sauvegarde le texte en fichier .txt"""
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
            return True
        except Exception as e:
            logger.error(f"Erreur lors de la sauvegarde en TXT: {e}")
            return False

def allowed_file(filename):
    """Vérifie si le fichier est autorisé"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ['doc', 'docx']

@app.route('/')
def index():
    """Page d'accueil de l'API"""
    return jsonify({
        'message': 'API de traitement de documents Word',
        'version': '1.0',
        'endpoints': {
            '/preview': 'POST - Aperçu du contenu du document',
            '/export': 'POST - Export du document dans différents formats',
            '/convert-to-html': 'POST - Convertir un document Word en HTML',
            '/export-to-word': 'POST - Exporter du HTML en document Word'
        }
    })

@app.route('/convert-to-html', methods=['POST'])
def convert_to_html():
    """Endpoint pour convertir un document Word en HTML"""
    try:
        # Vérifier si le fichier est présent dans la requête
        if 'file' not in request.files:
            return jsonify({'error': 'Aucun fichier fourni'}), 400
            
        file = request.files['file']
        
        # Vérifier si le fichier est autorisé
        if file.filename == '':
            return jsonify({'error': 'Aucun fichier sélectionné'}), 400
            
        if not file.filename.lower().endswith(('.doc', '.docx')):
            return jsonify({'error': 'Seuls les fichiers Word (.doc, .docx) sont acceptés'}), 400
        
        # Créer un fichier temporaire
        temp_dir = tempfile.mkdtemp()
        input_path = os.path.join(temp_dir, file.filename)
        
        # Sauvegarder le fichier téléchargé
        file.save(input_path)
        
        # Convertir en HTML
        html_content = DocumentProcessor.convert_docx_to_html(input_path)
        
        # Nettoyer le répertoire temporaire
        if os.path.exists(temp_dir):
            for f in os.listdir(temp_dir):
                os.remove(os.path.join(temp_dir, f))
            os.rmdir(temp_dir)
        
        return jsonify({
            'success': True,
            'html': html_content
        })
        
    except Exception as e:
        app.logger.error(f"Erreur lors de la conversion en HTML: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/export-to-word', methods=['POST'])
def export_to_word():
    """Endpoint pour exporter du HTML en document Word"""
    try:
        data = request.get_json()
        
        if not data or 'html' not in data:
            return jsonify({'error': 'Aucun contenu HTML fourni'}), 400
            
        html_content = data['html']
        file_name = data.get('fileName', 'document') + '.docx'
        
        # Créer un fichier temporaire
        temp_dir = tempfile.mkdtemp()
        output_path = os.path.join(temp_dir, file_name)
        
        # Convertir en DOCX
        DocumentProcessor.convert_html_to_docx(html_content, output_path)
        
        # Envoyer le fichier généré
        response = send_file(
            output_path,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=file_name
        )
        
        # Nettoyer le répertoire temporaire après l'envoi
        @response.call_on_close
        def remove_temp_files():
            try:
                if os.path.exists(output_path):
                    os.remove(output_path)
                if os.path.exists(temp_dir):
                    os.rmdir(temp_dir)
            except Exception as e:
                app.logger.error(f"Erreur lors du nettoyage des fichiers temporaires: {str(e)}")
        
        return response
        
    except Exception as e:
        app.logger.error(f"Erreur lors de l'export en Word: {str(e)}")
        return jsonify({'error': str(e)}), 500

# Redirection pour les anciennes routes
@app.route('/index.html')
def redirect_home():
    return redirect(url_for('home'))

@app.route('/editor.html')
def redirect_editor():
    return redirect(url_for('editor'))

if __name__ == '__main__':
    print("🚀 Démarrage du serveur Flask...")
    print("📄 Éditeur de documents Word")
    print(f"🌐 Application disponible sur: http://{Config.HOST}:{Config.PORT}")
    print("\nRoutes disponibles:")
    print(f"  - GET  /                : Page d'accueil")
    print(f"  - GET  /editor          : Éditeur de documents")
    print(f"  - POST /convert-to-html  : Convertir un document Word en HTML")
    print(f"  - POST /export-to-word   : Exporter du HTML en document Word")
    print("\nAppuyez sur Ctrl+C pour arrêter le serveur\n")
    
    # Créer le dossier static s'il n'existe pas
    os.makedirs('static', exist_ok=True)
    
    # Démarrer le serveur
    app.run(host=Config.HOST, port=Config.PORT, debug=Config.DEBUG)